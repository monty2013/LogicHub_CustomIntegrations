"""
name: MS Word
description: This integration is designed to create a word document partially based on a template with merge fields. It has the ability to adding pieces into the document.
logoUrl: https://upload.wikimedia.org/wikipedia/commons/8/88/MS_word_DOC_icon.svg
"""

from mailmerge import MailMerge
from docx import Document
from docx.shared import Inches
import json
import ast
from lhub_integ.params import ConnectionParam, ActionParam, InputType, JinjaTemplatedStr, DataType
from lhub_integ import action

STYLE = ActionParam("STYLE", description="Pick a style, refer to Word table styles for more details", input_type=InputType.SELECT, options=['Light List Accent 1','Light Grid Accent 1','Light Shading Accent 1'], default='Light List Accent 1', action="AppendTable")

base_path = "/opt/files/shared/integrationsFiles/"
events_base = "/opt/files/service/event_files/"

@action(name="Create New File")
def CreateNewFile(DocFile, Title):
    """
    This action allows you to create a new Document with a title. The document name will be returned for reference.
    :param DocFile: The column contains file name or ='docname.docx'.
    :param Title: The column contains the title or ='whatever title'
    :return: An instance to the file to be manipulated.
    """
    document = Document()
    document.add_heading(Title, 0)
    document.save(base_path + DocFile)
    return {
        "has_error": "false", "file_name": DocFile
    }


@action(name="Append Image File")
def AppendImage(DocFile, ImageFile):
    """
    The action allows you to insert an image file into the word document.
    :param DocFile: The column contains file name.
    :param ImageFile: The column contains Image file_id
    :return: An instance to the file to be manipulated.
    """
    document = Document(base_path+DocFile)
    document.add_picture(base_path+ImageFile, width=Inches(6.25))
    document.save(base_path+DocFile)
    return {
        "has_error":"false", "file_name":DocFile
    }

@action(name="Append Text")
def AppendText(DocFile, Paragraph:JinjaTemplatedStr):
    """
    The action allows you to append a paragraph into the word document.
    :param DocFile: The column contains file name.
    :param Paragraph: The column contains Paragraph
    :return: An instance to the file to be manipulated.
    """
    document = Document(base_path+DocFile)
    document.add_paragraph(Paragraph)
    document.save(base_path+DocFile)
    return {
        "has_error":"false", "file_name":DocFile
    }


@action(name="Append Heading")
def AppendHeading(DocFile, Heading):
    """
    The action allows you to append a new heading into the word document.
    :param DocFile: The column contains file name.
    :param Heading: The column contains the value or ="section heading"
    :return: An instance to the file to be manipulated.
    """
    document = Document(base_path+DocFile)
    document.add_heading(Heading)
    document.save(base_path+DocFile)
    return {
        "has_error":"false", "file_name":DocFile
    }

@action(name="Append Table")
def AppendTable(DocFile, Columns:int, Headers, RowArray: JinjaTemplatedStr):
    """
    The action allows you to append a matrix table onto the word document.
    :param DocFile: The column contains file name.
    :param Columns: The column specifies the number of columns or just enter =2 for 2 columns.
    :param Headers: comma separated column headers either from a parent column, or just ="col1,col2,etc"
    :param RowArray: The column contains Row data, the data structure should be a matrix, for example: [['a','b'],['c','d']], where rows and columns.
    :return: An instance to the file to be manipulated.
    """
    document = Document(base_path+DocFile)
    table = document.add_table(rows=1, cols=Columns)
    table.style=STYLE.read()
    headers = Headers.split(',')
    row = table.rows[0]
    col = 0
    for header in headers :
        row.cells[col].text = header
        col = col + 1
    data = ast.literal_eval(RowArray)
    for r in data :
        row = table.add_row()
        col = 0
        for c in r :
            row.cells[col].text = c
            col = col + 1
    document.save(base_path+DocFile)
    return {
        "has_error":"false", "file_name":DocFile
    }
    
@action(name="Apply Merge Fields")
def ApplyMergeFields(DocFile, NewFile, MergeJSON):
    """
    This action allows you to populate a word template file with MergeFields with a JSON data structure. The template file must be preloaded into the system under the service/event_files directory.
    :param DocFile: The column contains file name.
    :param NewFIle: The column contains new file name.
    :param MergeJSON: The column contains the JSON dictionary of the merge fields. This should match the template.
    :return: An instance to the file to be manipulated.
    """
    document = MailMerge(events_base+DocFile)
    mergeFields = json.loads(MergeJSON)
    document.merge_pages([mergeFields])
    document.write(base_path+NewFile)
    return {
        "has_error":"false", "file_name":NewFile
    }    
